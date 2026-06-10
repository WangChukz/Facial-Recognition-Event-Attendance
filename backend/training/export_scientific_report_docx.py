from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

TRAINING_DIR = Path(__file__).resolve().parent
OUT_DOCX = TRAINING_DIR / "results" / "BAO_CAO_CHI_TIET_KET_QUA_THUC_NGHIEM.docx"


def set_cell_shading(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = text
        set_cell_shading(cell, "1F4E79")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = str(text)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in [1, 2] and not text.replace(".", "").replace("%", "").strip().isdigit() else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    # Thêm khoảng trống sau bảng
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 13, "1F4E79"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def build_report() -> Document:
    doc = Document()
    setup_document(doc)

    # 4.3. Kết quả thực nghiệm
    h1 = doc.add_heading("4.3. Kết quả thực nghiệm và Thảo luận (Experimental Results)", level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)

    add_paragraph(
        doc,
        "Để đánh giá khả năng nhận diện và độ ổn định thực tế của hệ thống, chúng tôi tiến hành thực nghiệm "
        "và đối chiếu hiệu năng qua 5 thử nghiệm cụ thể dưới đây. Dữ liệu thử nghiệm bao gồm ảnh thực tế của "
        "39 sinh viên đi học, ảnh người lạ thu thập trên mạng, và bộ dữ liệu giả lập 16.000 người để thử nghiệm giới hạn tốc độ."
    )

    doc.add_heading("4.3.1. Các chỉ số đánh giá hiệu năng", level=2)
    add_paragraph(
        doc,
        "Chất lượng của hệ thống được đo lường cụ thể thông qua 07 chỉ số chính sau đây:"
    )

    # Bảng 4.7
    add_table(
        doc,
        ["STT", "Đo độ đo", "Mô tả ý nghĩa và tác dụng", "Đơn vị", "Kỳ vọng hệ thống"],
        [
            ["1", "Accuracy (Độ chính xác)", "Tỷ lệ nhận diện đúng danh tính trên tổng số ảnh sinh viên có trong hệ thống.", "%", "Càng cao càng tốt"],
            ["2", "Precision (Độ chuẩn xác)", "Trong những lần hệ thống báo nhận ra sinh viên, có bao nhiêu phần trăm là đúng.", "%", "Càng cao càng tốt"],
            ["3", "Recall (Độ nhạy)", "Tỷ lệ sinh viên được ghi nhận thành công trên tổng số sinh viên thực tế đứng trước camera.", "%", "Càng cao càng tốt"],
            ["4", "F1-Score (Chỉ số F1)", "Điểm đánh giá cân bằng tổng hợp giữa độ chuẩn xác (Precision) và độ nhạy (Recall).", "%", "Càng cao càng tốt"],
            ["5", "Similarity TB (Độ tương đồng)", "Điểm số khớp nhau trung bình giữa khuôn mặt camera quét được với ảnh thẻ gốc (từ 0 đến 1).", "0 – 1", "Càng cao càng tốt"],
            ["6", "Unknown Rejection (Lọc người lạ)", "Tỷ lệ từ chối thành công người ngoài (người lạ không có tên trong danh sách).", "%", "Càng cao càng tốt"],
            ["7", "Latency (Độ trễ xử lý)", "Thời gian hệ thống tìm kiếm và đối chiếu khuôn mặt (không tính thời gian phát hiện mặt).", "ms/frame", "Càng thấp càng tốt"]
        ],
        col_widths=[1.0, 3.5, 7.5, 2.0, 3.0]
    )

    # 4.3.2 Case 1
    doc.add_heading("4.3.2. Case 1: Đánh giá hiệu năng mô hình tinh chỉnh (Fine-tune ResNet-18)", level=2)
    add_paragraph(
        doc,
        "Ở thử nghiệm này, chúng tôi huấn luyện lại mô hình học sâu (ResNet-18) "
        "trên tập dữ liệu của 39 sinh viên (mỗi sinh viên chỉ có 1 ảnh gốc và được tạo thêm các ảnh biến thể để học). "
        "Mục tiêu là kiểm tra xem mô hình tự huấn luyện trên dữ liệu ít này có hoạt động tốt khi đưa vào lớp học thực tế hay không."
    )

    # Bảng kết quả Case 1
    add_table(
        doc,
        ["STT", "Thuật toán so khớp", "Accuracy", "Precision", "Recall", "F1-Score", "Similarity TB", "Unknown Rejection", "Latency Head"],
        [
            ["1", "FAISS Flat", "20.69%", "17.37%", "20.69%", "17.20%", "0.6425", "N/A", "0.0211 ms"],
            ["2", "FAISS HNSW", "20.69%", "19.53%", "20.69%", "17.15%", "0.6416", "N/A", "0.0080 ms"]
        ],
        col_widths=[1.0, 3.5, 2.0, 2.0, 2.0, 2.0, 1.8, 1.8, 2.5]
    )

    add_paragraph(
        doc,
        "Giải thích kết quả dễ hiểu: Độ chính xác của mô hình tự huấn luyện (Fine-tune) rất thấp, chỉ đạt 20.69% "
        "(tức là 10 lần điểm danh thì chỉ đúng khoảng 2 lần). Nguyên nhân là do mỗi sinh viên chỉ có 1 ảnh đăng ký ban đầu, "
        "khiến mô hình bị học vẹt (quá khớp). Nó chỉ nhận ra sinh viên khi đứng thẳng, đủ sáng giống hệt ảnh đăng ký. "
        "Khi áp dụng vào lớp học thực tế với camera mờ, ánh sáng thay đổi hoặc sinh viên hơi nghiêng đầu, mô hình hoàn toàn "
        "không nhận diện được. Kết luận là việc tự huấn luyện lại mô hình nhỏ trên dữ liệu ít là không khả thi."
    )

    # 4.3.3 Case 2
    doc.add_heading("4.3.3. Case 2: Đánh giá hiệu năng mô hình tiền huấn luyện (ArcFace Pretrained)", level=2)
    add_paragraph(
        doc,
        "Thử nghiệm này sử dụng trực tiếp mô hình ArcFace ResNet-50 đã được các chuyên gia huấn luyện sẵn trên hàng triệu khuôn mặt trước đó. "
        "Hệ thống rút trích đặc trưng của ảnh đăng ký gốc và so khớp trực tiếp với ảnh webcam phòng học thật của 39 sinh viên."
    )

    # Bảng kết quả Case 2
    add_table(
        doc,
        ["STT", "Thuật toán so khớp", "Accuracy", "Precision", "Recall", "F1-Score", "Similarity TB", "Unknown Rejection", "Latency Head"],
        [
            ["1", "FAISS Flat", "100.00%", "100.00%", "100.00%", "100.00%", "0.5967", "N/A", "0.0019 ms"],
            ["2", "FAISS HNSW", "100.00%", "100.00%", "100.00%", "100.00%", "0.5967", "N/A", "0.0031 ms"]
        ],
        col_widths=[1.0, 3.5, 2.0, 2.0, 2.0, 2.0, 1.8, 1.8, 2.5]
    )

    add_paragraph(
        doc,
        "Giải thích kết quả dễ hiểu: Mô hình dùng sẵn (Pretrained ArcFace) đạt độ chính xác tuyệt đối 100% trên tất cả các chỉ số "
        "(toàn bộ sinh viên đi học đều được điểm danh đúng). Sở dĩ mô hình này thông minh vượt trội là vì nó đã được học trước từ "
        "hàng triệu khuôn mặt khác nhau. Nhờ đó, nó tự rút ra được các đặc điểm cốt lõi trên khuôn mặt người (như khoảng cách mắt, mũi, miệng) "
        "mà không bị ảnh hưởng bởi môi trường. Dù ánh sáng lớp học có thay đổi hay sinh viên quay các góc mặt khác nhau, mô hình vẫn nhận diện chính xác."
    )

    # 4.3.4 Case 3
    doc.add_heading("4.3.4. Case 3: Thử nghiệm khả năng chịu tải và độ trễ ở quy mô lớn (N = 16.000)", level=2)
    add_paragraph(
        doc,
        "Khi số lượng sinh viên tăng lên hàng nghìn người, tốc độ tìm kiếm là yếu tố cực kỳ quan trọng để camera không bị giật lag. "
        "Chúng tôi thử nghiệm đo độ trễ tìm kiếm của thuật toán Flat (dò tìm tuần tự toàn bộ) và HNSW (dò tìm phân cấp thông minh) "
        "ở các quy mô từ 500 đến 16.000 sinh viên:"
    )

    # Bảng kết quả Case 3
    add_table(
        doc,
        ["Quy mô N (SV)", "Thuật toán so khớp", "Độ trễ tối thiểu (Min)", "Độ trễ tối đa (Max)", "Độ trễ trung bình (Mean)"],
        [
            ["N = 500", "FAISS Flat", "0.0051 ms", "0.0381 ms", "0.0089 ms"],
            ["(8.500 ảnh)", "FAISS HNSW", "0.0081 ms", "0.0401 ms", "0.0101 ms"],
            ["N = 1.000", "FAISS Flat", "0.0039 ms", "0.0456 ms", "0.0081 ms"],
            ["(17.000 ảnh)", "FAISS HNSW", "0.0076 ms", "0.0482 ms", "0.0108 ms"],
            ["N = 5.000", "FAISS Flat", "0.0121 ms", "0.1245 ms", "0.0163 ms"],
            ["(85.000 ảnh)", "FAISS HNSW", "0.0118 ms", "0.0651 ms", "0.0178 ms"],
            ["N = 16.000", "FAISS Flat", "0.0312 ms", "0.3840 ms", "0.0520 ms"],
            ["(272.000 ảnh)", "FAISS HNSW", "0.0175 ms", "0.1190 ms", "0.0321 ms"]
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5, 3.5]
    )

    add_paragraph(
        doc,
        "Giải thích kết quả dễ hiểu: Tốc độ tìm kiếm khuôn mặt phụ thuộc rất lớn vào cách tìm kiếm trong bộ nhớ:\n"
        "- Thuật toán Flat giống như việc ta đi dò tên từng người một từ đầu đến cuối danh sách. Khi lớp học chỉ có dưới 1.000 sinh viên, cách này vẫn chạy rất nhanh. Nhưng khi trường học có tới 16.000 sinh viên (tương đương 272.000 ảnh đặc trưng), việc dò từng người sẽ rất tốn thời gian, gây lag camera.\n"
        "- Thuật toán HNSW giống như một cuốn danh bạ thông minh được chia theo mục lục phân cấp. Thay vì dò hết, HNSW nhảy nhanh qua các nhóm khuôn mặt tương tự để tìm ra kết quả gần đúng nhất. Kết quả cho thấy ở quy mô 16.000 sinh viên, HNSW chỉ mất 0.0321 mili-giây để tìm ra khuôn mặt (nhanh hơn Flat 1.6 lần) và độ trễ tối đa cực kỳ thấp (0.1190 mili-giây). Do đó, HNSW là lựa chọn bắt buộc khi triển khai quy mô lớn để hệ thống chạy mượt mà."
    )

    # 4.3.5 Case 4
    doc.add_heading("4.3.5. Case 4: Đánh giá khả năng phát hiện người lạ (Unknown Rejection)", level=2)
    add_paragraph(
        doc,
        "Hệ thống sử dụng ngưỡng so khớp T = 0.45 làm mốc để nhận diện sinh viên. Nếu điểm tương đồng khuôn mặt nhỏ hơn 0.45, hệ thống sẽ coi là người ngoài lớp học và báo đỏ. "
        "Chúng tôi cho 85 ảnh người lạ đi qua camera để kiểm tra độ an toàn của ngưỡng quyết định này:"
    )

    # Bảng kết quả Case 4
    add_table(
        doc,
        ["STT", "Thuật toán so khớp", "Accuracy", "Precision", "Recall", "F1-Score", "Similarity TB", "Unknown Rejection", "Latency Head"],
        [
            ["1", "FAISS Flat", "N/A", "N/A", "N/A", "N/A", "0.1719", "100.00%", "0.0031 ms"],
            ["2", "FAISS HNSW", "N/A", "N/A", "N/A", "N/A", "0.1719", "100.00%", "0.0040 ms"]
        ],
        col_widths=[1.0, 3.5, 2.0, 2.0, 2.0, 2.0, 1.8, 1.8, 2.5]
    )

    add_paragraph(
        doc,
        "Giải thích kết quả dễ hiểu: Hệ thống đạt tỷ lệ lọc người lạ hoàn hảo 100.00%. Điểm tương đồng trung bình của người lạ rất thấp, chỉ đạt 0.1719, "
        "nằm dưới rất sâu so với ngưỡng quyết định 0.45. Điều này chứng minh hệ thống hoạt động vô cùng an toàn, tránh tối đa việc nhận diện nhầm "
        "người lạ hoặc điểm danh hộ cho người ngoài lớp học."
    )

    # 4.3.6 Case 5
    doc.add_heading("4.3.6. Case 5: Đánh giá cơ chế tự động tối ưu hóa ảnh mẫu (Progressive Gallery Enrichment)", level=2)
    add_paragraph(
        doc,
        "Progressive Gallery Enrichment là tính năng tự động lưu thêm ảnh thực tế của sinh viên khi họ điểm danh thành công với độ tin cậy cao (Similarity >= 0.75). "
        "Hệ thống áp dụng các bộ lọc bảo vệ nghiêm ngặt: chỉ lưu khi hệ thống cực kỳ tự tin (>=0.75), giới hạn số ảnh phụ được lưu, và giữ ảnh thẻ gốc làm neo chính "
        "để cơ sở dữ liệu không bao giờ bị lệch hướng (Gallery Drift)."
    )
    add_paragraph(
        doc,
        "Chúng tôi chia nhỏ dữ liệu thực tế (ảnh 1-3 để hệ thống tự động lưu mẫu phụ; ảnh 4-5 làm ảnh mới điểm danh) để đối chiếu hiệu quả:"
    )

    # Bảng kết quả Case 5
    add_table(
        doc,
        ["STT", "Trạng thái thư viện ảnh", "Accuracy", "Precision", "Recall", "F1-Score", "Similarity TB", "Unknown Rejection", "Latency Head"],
        [
            ["1", "Không cập nhật (Chỉ có ảnh thẻ)", "100.00%", "100.00%", "100.00%", "100.00%", "0.5782", "N/A", "0.0074 ms"],
            ["2", "Có cập nhật (Có ảnh phụ 1-3)", "100.00%", "100.00%", "100.00%", "100.00%", "0.7771", "N/A", "0.0057 ms"],
            ["3", "Có cập nhật + Người lạ (Lọc người ngoài)", "N/A", "N/A", "N/A", "N/A", "0.1938", "100.00%", "0.0066 ms"]
        ],
        col_widths=[1.0, 3.5, 2.0, 2.0, 2.0, 2.0, 1.8, 1.8, 2.5]
    )

    add_paragraph(
        doc,
        "Giải thích kết quả dễ hiểu: Khi bật tính năng tự động cập nhật ảnh mẫu, độ tương đồng trung bình khi nhận diện sinh viên tăng vọt từ 0.5782 lên 0.7771 (tăng +19.89%). "
        "Điều này giúp hệ thống 'quen mặt' sinh viên hơn sau mỗi buổi đi học, giúp những lần điểm danh tiếp theo diễn ra nhanh nhạy hơn, dễ dàng vượt qua các biến đổi về ánh sáng hoặc góc nghiêng. "
        "Đồng thời, tỷ lệ từ chối người ngoài vẫn giữ vững ở mức tuyệt đối 100.00% với Similarity trung bình của người lạ thấp (0.1938), "
        "chứng minh cơ chế tự động cập nhật này hoạt động vô cùng an toàn và không gây lỗi nhầm lẫn."
    )

    return doc


def main() -> None:
    doc = build_report()
    os.makedirs(OUT_DOCX.parent, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Successfully exported cleaned-up scientific report to: {OUT_DOCX}")


if __name__ == "__main__":
    main()
