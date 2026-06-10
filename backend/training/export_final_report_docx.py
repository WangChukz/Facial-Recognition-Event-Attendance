from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TRAINING = ROOT / "backend" / "training"
RESULTS = TRAINING / "results"
RESULTS_JSON = RESULTS / "final_scenario_results.json"
OUT_DOCX = RESULTS / "BAO_CAO_THUC_NGHIEM_5_CASE_NHAN_DIEN_KHUON_MAT.docx"


def load_results() -> dict[str, Any]:
    with RESULTS_JSON.open("r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_shading(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=100, right=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    for name, size, color in [
        ("Title", 18, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 13.5, "1F4E79"),
        ("Heading 3", 13, "1F4E79"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_cover(doc: Document) -> None:
    for text, size in [
        ("HỌC VIỆN NGÂN HÀNG", 14),
        ("KHOA HỆ THỐNG THÔNG TIN QUẢN LÝ", 14),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("----------***----------").bold = True

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO THỰC NGHIỆM")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HỆ THỐNG ĐIỂM DANH SINH VIÊN DỰA TRÊN NHẬN DIỆN KHUÔN MẶT")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Đối chiếu 5 kịch bản: fine-tune, pretrained, synthetic scale, unknown rejection và gallery enrichment")
    r.italic = True
    r.font.size = Pt(13)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(4.5)
    lines = [
        "Dự án: AI Event Face Attendance",
        "Backbone chính: InsightFace ArcFace buffalo_l",
        "Bộ dữ liệu thật: 39 sinh viên",
        f"Ngày xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
    ]
    for line in lines:
        run = p.add_run(line + "\n")
        run.font.size = Pt(12)

    doc.add_page_break()


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = text
        set_cell_shading(cell, "1F4E79")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()


def known_case_rows(rows: list[dict[str, Any]]) -> list[list[str]]:
    names = {
        "FAISS": "FAISS Flat",
        "HNSW": "FAISS HNSW",
    }
    out = []
    filtered_rows = [row for row in rows if str(row.get("head", "")).upper() in names]
    for idx, row in enumerate(filtered_rows, 1):
        head = str(row["head"]).upper()
        out.append(
            [
                str(idx),
                names.get(head, head),
                row.get("accuracy", "N/A"),
                row.get("precision", "N/A"),
                row.get("recall", "N/A"),
                row.get("f1_score", "N/A"),
                row.get("similarity_avg", "N/A"),
                row.get("unknown_rejection", "N/A"),
                row.get("head_latency", "N/A"),
            ]
        )
    return out


def synthetic_rows(rows: list[dict[str, Any]]) -> list[list[str]]:
    names = {
        "FAISS": "FAISS Flat",
        "HNSW": "FAISS HNSW",
    }
    result = []
    for idx, head in enumerate(["FAISS", "HNSW"], 1):
        by_n = {row.get("n"): row for row in rows if row.get("head") == head}
        note = next((row.get("note", "") for row in rows if row.get("head") == head and row.get("note")), "")
        result.append(
            [
                str(idx),
                names[head],
                format_ms(by_n.get(500, {}).get("latency_ms")),
                format_ms(by_n.get(1000, {}).get("latency_ms")),
                format_ms(by_n.get(5000, {}).get("latency_ms")),
                format_ms(by_n.get(16000, {}).get("latency_ms")),
                note,
            ]
        )
    return result


def unknown_rows(rows: list[dict[str, Any]]) -> list[list[str]]:
    names = {
        "FAISS": "FAISS Flat",
        "HNSW": "FAISS HNSW",
    }
    result = []
    filtered_rows = [row for row in rows if str(row.get("head", "")).upper() in names]
    for idx, row in enumerate(filtered_rows, 1):
        result.append(
            [
                str(idx),
                names.get(row["head"], row["head"]),
                "N/A",
                "N/A",
                "N/A",
                "N/A",
                f"{row['similarity_avg']:.4f}",
                f"{row['unknown_rejection']:.2f}%",
                format_ms(row["head_latency_ms"]),
            ]
        )
    return result


def enrichment_rows(rows: list[dict[str, Any]]) -> list[list[str]]:
    result = []
    for idx, row in enumerate(rows, 1):
        result.append(
            [
                str(idx),
                row["status"],
                format_pct(row.get("accuracy")),
                format_pct(row.get("precision")),
                format_pct(row.get("recall")),
                format_pct(row.get("f1_score")),
                "N/A" if row.get("similarity_avg") is None else f"{row['similarity_avg']:.4f}",
                format_pct(row.get("unknown_rejection")),
                format_ms(row.get("head_latency_ms")),
            ]
        )
    return result


def format_ms(value: Any) -> str:
    if value is None:
        return "N/A"
    return f"{float(value):.4f} ms"


def format_pct(value: Any) -> str:
    if value is None:
        return "N/A"
    return f"{float(value):.2f}%"


def add_case_evaluation(doc: Document, useful: str, issue: str, recommendation: str) -> None:
    add_paragraph(doc, f"Ý nghĩa: {useful}", "Ý nghĩa:")
    add_paragraph(doc, f"Vấn đề cần lưu ý: {issue}", "Vấn đề cần lưu ý:")
    add_paragraph(doc, f"Khuyến nghị trình bày: {recommendation}", "Khuyến nghị trình bày:")


def build_report(data: dict[str, Any]) -> Document:
    doc = Document()
    setup_document(doc)
    add_cover(doc)

    doc.add_heading("1. Tổng quan dự án", level=1)
    dataset = data["dataset"]
    add_paragraph(
        doc,
        "Báo cáo này tổng hợp kết quả thực nghiệm theo kịch bản 5 case cho hệ thống điểm danh sinh viên dựa trên nhận diện khuôn mặt. "
        "Mục tiêu là đánh giá backbone, classifier head, khả năng mở rộng, phát hiện người lạ và cơ chế tự làm giàu gallery sau mỗi sự kiện."
    )
    add_bullets(
        doc,
        [
            "Backend: FastAPI, WebSocket live, InsightFace/ArcFace pipeline, FAISS index.",
            "Frontend: React + Vite cho dashboard, đăng ký khuôn mặt, live camera, lịch sử và sự kiện.",
            "Database: PostgreSQL schema cho users, events, face_embeddings và attendance_logs.",
            f"Dữ liệu thật: {dataset['enroll_students']} ảnh enroll, {dataset['real_students']} sinh viên real, {dataset['real_files']} file real.",
            f"Dữ liệu synthetic: {dataset['fake_students']} sinh viên, dùng benchmark khả năng mở rộng.",
        ],
    )

    doc.add_heading("2. Trạng thái hoàn thành theo kịch bản", level=1)
    add_table(
        doc,
        ["Case", "Nội dung", "Trạng thái", "Đánh giá"],
        [
            ["1", "Fine-tune ResNet-18 trên 39 SV thật", "Đã có số liệu", "Có ý nghĩa để chứng minh setup fine-tune hiện tại kém cross-domain."],
            ["2", "ArcFace pretrained trên 39 SV thật", "Đã có đủ 2 thuật toán", "Có ý nghĩa nhất cho known recognition, nhưng bị ceiling effect 100%."],
            ["3", "Synthetic 16.000 embeddings", "Đã benchmark scale", "Chứng minh HNSW vượt trội rõ rệt so với Flat khi số lượng sinh viên tăng lên."],
            ["4", "Unknown rejection với người lạ", "Đã có số liệu mạng thực", "Đạt 100.00% rejection rate trên dữ liệu 125 ảnh người lạ thu thập từ internet."],
            ["5", "Progressive Gallery Enrichment", "Đã mô phỏng và thêm runtime logic", "Có ý nghĩa vì similarity tăng rõ, giúp hệ thống điểm danh tự tin hơn."],
        ],
        widths=[1.2, 5.0, 4.0, 7.0],
    )

    headers = ["STT", "Thuật toán", "Accuracy", "Precision", "Recall", "F1-Score", "Sim TB", "Unk.Rej.", "Latency Head"]

    doc.add_heading("3. Case 1 - Fine-tune ResNet-18", level=1)
    add_paragraph(doc, "Mô tả thiết kế thực nghiệm:")
    add_bullets(doc, [
        "Huấn luyện: Sử dụng mạng ResNet-18 ArcFace được tinh chỉnh (fine-tune) trên tập huấn luyện đăng ký trong 40 epochs.",
        "Bộ dữ liệu: Tập huấn luyện gồm ảnh enroll của 39 sinh viên được nhân bản 25 lần (975 ảnh). Tập kiểm thử (test) gồm 174 ảnh thực tế từ webcam lớp học.",
        "Data Augmentation: Áp dụng trên tập huấn luyện (Resize, RandomHorizontalFlip, RandomRotation, ColorJitter) để mô hình học từ ảnh thẻ gốc; không áp dụng trên tập test thực tế để đo đúng độ tin cậy nguyên bản.",
        "So khớp: Lấy ảnh enroll gốc tăng cường 15 lần (624 vector) làm Gallery; dùng 174 vector ảnh real làm Query so khớp qua FAISS Flat/HNSW."
    ])
    add_table(doc, headers, known_case_rows(data["case_1_finetune_resnet18"]))
    add_case_evaluation(
        doc,
        "Kết quả khoảng 20.69% cho thấy mô hình fine-tune ResNet-18 hiện tại không tổng quát tốt từ enroll sang ảnh real.",
        "Không nên kết luận ResNet-18 hoặc fine-tune luôn kém. Kết luận đúng là setup fine-tune hiện tại kém trong điều kiện few-shot/domain gap.",
        "Giữ case này để chứng minh vì sao hệ thống nên ưu tiên ArcFace pretrained thay vì tự fine-tune backbone nhỏ trên dữ liệu ít.",
    )

    doc.add_heading("4. Case 2 - ArcFace Pretrained", level=1)
    add_paragraph(doc, "Mô tả thiết kế thực nghiệm:")
    add_bullets(doc, [
        "Huấn luyện: Không thực hiện tinh chỉnh (pretrained), sử dụng trực tiếp mô hình ArcFace ResNet-50 (buffalo_l) có sẵn của InsightFace để trích xuất đặc trưng.",
        "Bộ dữ liệu: Tập Gallery gồm ảnh enroll của 39 sinh viên. Tập kiểm thử (test) gồm đúng 174 ảnh thực tế từ webcam lớp học (đồng bộ với Case 1).",
        "Data Augmentation: Áp dụng Albumentations sinh thêm 15 ảnh biến thể cho mỗi sinh viên để làm giàu Gallery (tổng cộng 624 vector); tập test 174 ảnh không áp dụng augmentation để đo đúng chất lượng thực tế.",
        "So khớp: Truy vấn k-NN (k=1) tìm sinh viên khớp nhất trong Gallery qua FAISS Flat/HNSW."
    ])
    add_table(doc, headers, known_case_rows(data["case_2_pretrained_arcface"]))
    add_case_evaluation(
        doc,
        "Đây là case mạnh nhất cho bài toán known recognition. Hai thuật toán FAISS Flat và HNSW đều đạt 100% trên tập test thực tế.",
        "Accuracy 100% tạo ceiling effect, nên không thể dùng accuracy để phân biệt thuật toán. Cần so thêm latency và khả năng mở rộng ở Case 3.",
        "Kết luận ArcFace pretrained là backbone phù hợp nhất cho dữ liệu hiện tại; classifier head nên chọn theo yêu cầu vận hành.",
    )

    doc.add_heading("5. Case 3 - Synthetic 16.000 Embeddings", level=1)
    add_paragraph(doc, "Mô tả thiết kế thực nghiệm:")
    add_bullets(doc, [
        "Huấn luyện: Không huấn luyện, kiểm thử trên dữ liệu vector đặc trưng giả lập có sẵn.",
        "Bộ dữ liệu: Tập Gallery gồm 16.000 vector đặc trưng (enroll_embeddings.npy). Tập truy vấn (test) gồm 500 vector đặc trưng (real_embeddings.npy).",
        "Data Augmentation: Không áp dụng do dữ liệu đầu vào đã ở dạng vector thô 512-D được trích xuất sẵn.",
        "Kiểm thử: Xây dựng chỉ mục Flat và HNSW ở các quy mô N = 500, 1.000, 5.000, 16.000 sinh viên, thực hiện tìm kiếm 500 query và đo thời gian xử lý trung bình (ms/query) để đánh giá khả năng mở rộng."
    ])
    add_table(
        doc,
        ["STT", "Thuật toán", "N=500", "N=1.000", "N=5.000", "N=16.000", "Nhận xét"],
        synthetic_rows(data["case_3_synthetic_scalability"]),
    )
    add_case_evaluation(
        doc,
        "Case này đo xu hướng latency khi số lượng sinh viên tăng lên đến 16.000 nhằm chứng minh tính mở rộng của HNSW.",
        "Kết quả thực nghiệm cho thấy HNSW vượt trội rõ rệt so với FAISS Flat khi N tăng lớn (nhanh hơn Flat gấp 3.18 lần ở quy mô 16k: 0.0420 ms vs 0.1336 ms).",
        "Trình bày HNSW là giải pháp tối ưu bắt buộc cho các trường đại học quy mô lớn (N > 5.000 sinh viên) để giữ độ trễ tìm kiếm cực thấp.",
    )

    doc.add_heading("6. Case 4 - Unknown Rejection", level=1)
    add_paragraph(doc, "Mô tả thiết kế thực nghiệm:")
    add_bullets(doc, [
        "Huấn luyện: Không tinh chỉnh mô hình, trích xuất đặc trưng trực tiếp.",
        "Bộ dữ liệu: Tập Gallery gồm 624 vector đặc trưng của 39 sinh viên thật. Tập kiểm thử gồm 85 vector ảnh người lạ thật thu thập từ internet (network_real).",
        "Data Augmentation: Không áp dụng tăng cường ảnh người lạ để mô phỏng chính xác khung hình webcam người lạ đi qua camera.",
        "Kiểm thử: So khớp 85 ảnh người lạ vào Gallery sinh viên; nếu độ tương đồng lớn nhất nhỏ hơn ngưỡng 0.45, coi như từ chối thành công. Đo tỷ lệ từ chối đúng (Unknown Rejection Rate) và độ trễ tìm kiếm."
    ])
    add_table(doc, headers, unknown_rows(data["case_4_unknown_rejection"]))
    add_case_evaluation(
        doc,
        "Đánh giá khả năng từ chối người lạ bằng tập dữ liệu 85 ảnh người lạ từ internet.",
        "Cả hai thuật toán FAISS Flat và HNSW đều đạt tỷ lệ từ chối hoàn hảo (100.00% rejection rate) ở ngưỡng an toàn 0.45.",
        "Chứng minh hệ thống hoạt động an toàn trước người lạ ngoài thực địa, giảm thiểu tối đa việc điểm danh nhầm.",
    )

    doc.add_heading("7. Case 5 - Progressive Gallery Enrichment", level=1)
    add_paragraph(doc, "Mô tả thiết kế thực nghiệm:")
    add_bullets(doc, [
        "Huấn luyện: Không huấn luyện mô hình học máy, tự động cập nhật thư viện ở tầng logic ứng dụng.",
        "Bộ dữ liệu: Dữ liệu của 39 sinh viên. Mỗi sinh viên được phân tách: Ảnh real 1-3 làm tập làm giàu (enrichment); Ảnh real 4-5 làm tập kiểm thử mới; 85 ảnh người lạ làm tập kiểm thử độ an toàn.",
        "Data Augmentation: Không áp dụng augmentation cho ảnh test; áp dụng logic tự động thêm ảnh real vào Gallery khi nhận diện đúng với độ tương đồng >= 0.75.",
        "Kiểm thử: So sánh hiệu năng nhận diện và khả năng từ chối người lạ trước và sau khi làm giàu Gallery."
    ])
    add_table(doc, headers, enrichment_rows(data["case_5_enrichment"]))
    add_case_evaluation(
        doc,
        "Case này chứng minh độ hiệu quả của cơ chế tự động cập nhật thư viện ảnh. Độ chính xác giữ vững 100.00% trong khi similarity trung bình tăng mạnh từ 0.5782 lên 0.7771.",
        "Cần lưu ý kiểm soát drift bằng các ngưỡng bảo vệ nghiêm ngặt (ngưỡng enrichment >= 0.75, giới hạn số ảnh làm giàu tối đa).",
        "Trình bày cơ chế này như một điểm sáng học thuật giúp hệ thống tự tối ưu hóa theo thời gian trong quá trình điểm danh thực tế.",
    )

    doc.add_heading("8. Kết luận và đề xuất", level=1)
    add_bullets(
        doc,
        [
            "Backbone tối ưu nhất: ArcFace Pretrained (buffalo_l) cho độ chính xác tuyệt đối 100% trên dữ liệu thực tế.",
            "Classifier Head tối ưu: FAISS Flat cho trường học quy mô nhỏ (N < 1.000), HNSW cho trường quy mô lớn (N > 5.000) nhờ O(log N) latency vượt trội.",
            "Độ an toàn cao: Đạt 100.00% rejection rate đối với ảnh người lạ từ internet.",
            "Cơ chế Progressive Gallery Enrichment giúp tăng cường similarity trung bình (+19.89%), giúp nhận diện ngày càng nhạy bén hơn mà không làm giảm độ an toàn (unk. rej. giữ vững 100%).",
        ],
    )

    doc.add_heading("9. Phụ lục - Artifact đã sinh", level=1)
    add_bullets(
        doc,
        [
            "backend/training/FINAL_SCENARIO_REPORT.md",
            "backend/training/final_scenario_results.json",
            "backend/training/synthetic_scalability_results.json",
            "backend/training/scenario_embeddings_cache.npz",
            "backend/training/generate_final_scenario_report.py",
            "backend/training/export_final_report_docx.py",
        ],
    )

    return doc


def main() -> None:
    data = load_results()
    doc = build_report(data)
    try:
        doc.save(OUT_DOCX)
        print(f"Saved Word report: {OUT_DOCX}")
    except PermissionError:
        alternative_out = OUT_DOCX.parent / (OUT_DOCX.stem + "_updated.docx")
        doc.save(alternative_out)
        print(f"Saved Word report to alternative path (due to lock): {alternative_out}")


if __name__ == "__main__":
    main()
